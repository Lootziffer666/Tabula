from __future__ import annotations

import difflib
from pathlib import Path

from docx import Document

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None


def smart_merge_documents(file1: Path, file2: Path) -> tuple[Path, str]:
    ext = file1.suffix.lower()
    master_path = file1.parent / f"MERGED_{file1.stem}_PLUS_{file2.stem}{ext if ext in {'.txt', '.docx'} else '.txt'}"

    if ext == ".txt":
        text1 = file1.read_text(encoding="utf-8", errors="ignore")
        text2 = file2.read_text(encoding="utf-8", errors="ignore")
        diff = list(difflib.unified_diff(text1.splitlines(), text2.splitlines(), lineterm=""))
        merged = text1
        for line in diff:
            if line.startswith("+") and not line.startswith("+++"):
                merged += "\n[ÄNDERUNG aus Datei 2] " + line[1:]
        master_path.write_text(merged, encoding="utf-8")
        return master_path, "TXT-Merge fertig"

    if ext == ".docx":
        doc1 = Document(file1)
        doc2 = Document(file2)
        merged_doc = Document()
        for para in doc1.paragraphs:
            merged_doc.add_paragraph(para.text)
        merged_doc.add_heading("ÄNDERUNGEN AUS DATEI 2", level=2)
        for para in doc2.paragraphs:
            p = merged_doc.add_paragraph(para.text)
            p.add_run("  ← aus Datei 2").bold = True
        merged_doc.save(master_path)
        return master_path, "DOCX-Merge fertig"

    if ext == ".pdf" and fitz is not None:
        doc = fitz.open(file1)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        out = master_path.with_suffix(".txt")
        out.write_text(text + f"\n\n[MERGE aus {file2.name}]", encoding="utf-8")
        return out, "PDF-Textmerge fertig"

    # fallback
    out = master_path.with_suffix(".txt")
    out.write_text(file1.read_text(encoding="utf-8", errors="ignore"), encoding="utf-8")
    return out, "Fallback-Merge erstellt"
